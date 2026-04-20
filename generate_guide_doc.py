#!/usr/bin/env python3
"""生成 TARDS 卡牌效果编写指南 Word 文档。"""

import sys
sys.path.insert(0, r"C:\Users\34773\Desktop\新建文件夹\.venv\Lib\site-packages")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

OUTPUT_PATH = r"C:\Users\34773\Desktop\新建文件夹\TARDS卡牌效果编写指南.docx"


def set_chinese_font(run, font_name="微软雅黑", size=11, bold=False, color=None):
    """设置中文字体。"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = color


def add_heading_zh(doc, text, level=1):
    """添加中文标题。"""
    sizes = {1: 22, 2: 16, 3: 14}
    colors = {1: RGBColor(0x2E, 0x5E, 0xAA), 2: RGBColor(0x33, 0x33, 0x33), 3: RGBColor(0x44, 0x44, 0x44)}
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_chinese_font(run, "微软雅黑", sizes.get(level, 12), bold=True, color=colors.get(level))
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph_zh(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, indent=True):
    """添加中文段落。"""
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_chinese_font(run, "微软雅黑", size, bold=bold, color=color)
    return p


def add_code_block(doc, code_lines):
    """添加代码块（等宽字体，浅灰背景）。"""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(line)
        set_chinese_font(run, "Consolas", 10, bold=False, color=RGBColor(0x22, 0x22, 0x22))
    # 空行
    doc.add_paragraph()


def add_box(doc, title, lines):
    """添加一个带标题的信息框。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(f"【{title}】")
    set_chinese_font(run, "微软雅黑", 11, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(f"  • {line}")
        set_chinese_font(run, "微软雅黑", 10.5)
    doc.add_paragraph()


def add_checklist(doc, items):
    """添加自检清单（带复选框样式）。"""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(f"☐ {item}")
        set_chinese_font(run, "微软雅黑", 10.5)
    doc.add_paragraph()


def main():
    doc = Document()

    # ===================== 标题页 =====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TARDS 卡牌效果编写指南")
    set_chinese_font(run, "微软雅黑", 28, bold=True, color=RGBColor(0x2E, 0x5E, 0xAA))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("（队友协作版 · 人人能看懂）")
    set_chinese_font(run, "微软雅黑", 14, bold=False, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("版本：1.0    |    适用：underworld / discrete / blood 三个卡包")
    set_chinese_font(run, "微软雅黑", 11, color=RGBColor(0x88, 0x88, 0x88))

    doc.add_page_break()

    # ===================== 第一章：核心概念 =====================
    add_heading_zh(doc, "第一章  我们在做什么", level=1)

    add_paragraph_zh(doc,
        "每张卡牌都有『效果』。比如『松鼠球』的效果是：受伤后向旁边移动，并在原地留下一只松鼠。"
        "这些效果需要写成代码。问题是：我们四个人各自写，很容易格式不统一，甚至写出互相冲突的代码。"
        "本文档的目标只有一个：让四个人写出来的效果代码像一个人写的一样整齐。")

    add_heading_zh(doc, "1.1  三个核心比喻", level=2)

    add_paragraph_zh(doc, "为了让大家不迷失在代码里，先把整个系统用三个比喻讲清楚：", bold=True)

    add_box(doc, "比喻一：EffectUtils = 厨房里的标准菜刀",
        [
            "『把肉切块』这个操作，你不需要自己磨一把刀，直接拿厨房里的标准菜刀就行。",
            "EffectUtils 就是那些『标准菜刀』——把单位返回手牌、召唤 token、造成伤害……都是现成的。",
            "禁止自己『造菜刀』（手写重复逻辑），必须用厨房里的。"
        ])

    add_box(doc, "比喻二：@special 装饰器 = 上菜前的检查清单",
        [
            "你做好了一道菜，服务员端出去之前要检查：盘子有没有？盐放了没？温度对吗？",
            "@special 就是这道检查。游戏启动时会自动检查你的代码格式。",
            "如果缺了关键步骤，游戏直接报错告诉你，而不是端给玩家吃出毛病。"
        ])

    add_box(doc, "比喻三：SPECIAL_MAP = 菜单登记簿",
        [
            "餐厅有一本菜单登记簿，上面写着『松鼠球』对应『_songshuqiu_special』这道菜。",
            "你写好了代码，必须去登记簿上登记，否则顾客点了『松鼠球』，厨房不知道做哪道菜。"
        ])

    # ===================== 第二章：六步工作流程 =====================
    doc.add_page_break()
    add_heading_zh(doc, "第二章  完整工作流程（六步走）", level=1)

    add_paragraph_zh(doc,
        "拿到一张卡牌后，按照以下六个步骤操作。不要跳过任何一步。"
        "每一步都有明确的『输入』和『输出』，你只需要填空。")

    # 步骤 0
    add_heading_zh(doc, "步骤 0：准备工作（只需做一次）", level=2)
    add_paragraph_zh(doc, "确认你的电脑上有以下文件（项目根目录下）：")
    add_code_block(doc, [
        "card_pools/effect_utils.py      ← 标准工具库（厨房菜刀）",
        "card_pools/effect_decorator.py  ← 检查装饰器（上菜检查）",
        "tools/gen_effect.py             ← 骨架生成工具（自动模板）"
    ])
    add_paragraph_zh(doc, "如果缺少，找组长拷贝一份。")

    # 步骤 1
    add_heading_zh(doc, "步骤 1：判断触发时机", level=2)
    add_paragraph_zh(doc, "拿到卡牌后，先看效果描述里的关键词，判断它什么时候触发：")

    # 用表格展示触发时机
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '触发时机'
    hdr_cells[1].text = '关键词'
    hdr_cells[2].text = '例子'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, "微软雅黑", 11, bold=True)

    rows = [
        ("部署时", "部署时、部署：", "鹏：部署时让敌方单位回手牌"),
        ("回合结束时", "回合结束：", "天牛：回合结束时移动敌方单位"),
        ("受到伤害时", "受到伤害后", "松鼠球"),
        ("死亡时", "亡语：", "某卡：死亡时对敌方造成伤害"),
        ("回合开始时", "回合开始：", "某卡：回合开始时加攻击力"),
    ]
    for timing, keyword, example in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = timing
        row_cells[1].text = keyword
        row_cells[2].text = example
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, "微软雅黑", 10.5)

    doc.add_paragraph()
    add_paragraph_zh(doc, "判断好后记下来，下一步要用。")

    # 步骤 2
    add_heading_zh(doc, "步骤 2：生成代码骨架", level=2)
    add_paragraph_zh(doc, "打开命令行（Windows 按 Win+R，输入 cmd，回车），进入项目目录，运行以下命令：")
    add_code_block(doc, [
        "cd C:\\Users\\34773\\Desktop\\新建文件夹",
        "python tools/gen_effect.py \"松鼠球(铁)\" \"受到伤害后向相邻移动一格，在原地留下松鼠\" --type damage_taken --file underworld"
    ])
    add_paragraph_zh(doc, "命令格式拆解：")
    add_code_block(doc, [
        "python tools/gen_effect.py  \"卡牌名(稀有度)\"  \"效果描述\"  --type <触发时机>  --file <卡包名>"
    ])
    add_paragraph_zh(doc, "其中 --type 可选值：deploy（部署时）、turn_start（回合开始）、turn_end（回合结束）、damage_taken（受伤时）、deathrattle（亡语）。")
    add_paragraph_zh(doc, "运行后，命令行会输出一段代码。把它复制下来，粘贴到对应的 effects 文件里。例如 underworld 卡包的效果写在 card_pools/underworld_effects.py 中。")

    add_paragraph_zh(doc, "【输出示例】", bold=True)
    add_code_block(doc, [
        "# ===== 松鼠球(铁) =====",
        "# 效果描述：受到伤害后向相邻移动一格，在原地留下松鼠",
        "# 触发时机：damage_taken",
        "# 目标：自动触发",
        "# 状态变量：_songshuqiu_xxx（用完即删）",
        "def _songshuqiu_special(minion: \"Minion\", player: \"Player\", game: \"Game\", extras=None) -> None:",
        '    """TODO: 实现松鼠球(铁)的damage_taken效果。',
        "",
        '        效果描述：受到伤害后向相邻移动一格，在原地留下松鼠',
        '    """',
        '    # TODO 在此实现效果',
        '    print(f"  [TODO] {minion.name} 的 special_fn 尚未实现")'
    ])

    # 步骤 3
    add_heading_zh(doc, "步骤 3：填空实现（使用标准工具）", level=2)
    add_paragraph_zh(doc, "删掉 print([TODO]...) 那行，换成真正的效果代码。")
    add_paragraph_zh(doc, "【第一步】打开 effect_utils.py，查找有没有现成的工具可以用：", bold=True)
    add_code_block(doc, [
        "# effect_utils.py 里的常用工具",
        "",
        "return_minion_to_hand(minion, game)      # 把场上单位返回手牌（满则弃置）",
        "summon_token(game, name, owner, pos)     # 在指定位置召唤一个 token 单位",
        "get_adjacent_positions(pos, board)       # 获取上下左右四个相邻位置",
        "deal_damage_to_minion(target, dmg)       # 对单位造成伤害（自动触发护盾/坚韧）",
        "create_echo_card(source_card, level)     # 创建回响版本卡牌",
    ])

    add_paragraph_zh(doc, "【第二步】编写效果代码。以松鼠球为例：", bold=True)
    add_code_block(doc, [
        "from card_pools.effect_utils import get_adjacent_positions, summon_token",
        "import random",
        "",
        "def _songshuqiu_special(minion: \"Minion\", player: \"Player\", game: \"Game\", extras=None) -> None:",
        '    """松鼠球：受到伤害后向相邻移动一格，在原地留下松鼠。"""',
        "    def on_damage(m=minion, g=game, p=player):",
        "        if getattr(m, '_songshuqiu_triggered', False):",
        "            return",
        "        m._songshuqiu_triggered = True",
        "        old_pos = m.position",
        "        candidates = get_adjacent_positions(old_pos, g.board)",
        "        valid = [pos for pos in candidates if pos not in g.board.minion_place]",
        "        if not valid:",
        '            print(f"  {m.name} 受伤但无处可去")',
        "            return",
        "        new_pos = random.choice(valid)",
        "        g.move_minion(m, new_pos)",
        "        summon_token(game=g, name='松鼠', owner=p, position=old_pos, attack=1, health=1)",
        '        print(f"  {m.name} 受伤后移动至 {new_pos}，在 {old_pos} 留下松鼠")',
        "    m._on_take_combat_damage.append(on_damage)",
    ])

    add_paragraph_zh(doc, "【关键细节】回调函数里的 m=minion 写法很重要，这叫『默认参数绑定』。如果不这么写，游戏会指向错误的单位。照抄即可，不需要理解原理。", bold=True, color=RGBColor(0xC0, 0x00, 0x00))

    # 步骤 4
    add_heading_zh(doc, "步骤 4：加上装饰器（自动检查）", level=2)
    add_paragraph_zh(doc, "在函数定义前面加一行 @special，就像给信封贴邮票一样简单：")
    add_code_block(doc, [
        "from card_pools.effect_decorator import special",
        "",
        "@special                                    # ← 加上这一行",
        "def _songshuqiu_special(minion, player, game, extras=None):",
        "    ..."
    ])
    add_paragraph_zh(doc, "@special 会自动检查三件事：")
    add_box(doc, "装饰器检查内容",
        [
            "参数必须有 minion、player、game、extras 四个——少一个游戏启动直接报错",
            "extras 必须有默认值 =None——没写也会报错",
            "函数必须有文档字符串（\"\"\" ... \"\"\"）——没写会弹警告"
        ])

    # 步骤 5
    add_heading_zh(doc, "步骤 5：更新 SPECIAL_MAP", level=2)
    add_paragraph_zh(doc, "在 effects 文件最上方，找到 SPECIAL_MAP，添加一行登记：")
    add_code_block(doc, [
        "SPECIAL_MAP = {",
        '    "松鼠球(铁)": "_songshuqiu_special",     # ← 加上这一行',
        '    "雕(铁)": "_diao_special",',
        '    "鹏(铁)": "_peng_special",',
        "    ...",
        "}"
    ])
    add_paragraph_zh(doc, "左边是卡牌在游戏里的完整名称（包括括号里的稀有度），右边是你刚才写的函数名。必须完全一致，包括括号和汉字。")

    # 步骤 6
    add_heading_zh(doc, "步骤 6：自检并测试", level=2)
    add_paragraph_zh(doc, "对照以下清单，全部打勾后再提交：")
    add_checklist(doc, [
        "函数签名正确：def _xxx_special(minion, player, game, extras=None) -> None:",
        "加了 @special 装饰器",
        "有文档字符串，包含原始卡面描述",
        "状态变量以 _{卡牌缩写}_ 开头（如 _songshuqiu_triggered）",
        "使用了 effect_utils 里的标准工具（如果有的话）",
        "操作前检查了 minion.is_alive()",
        "每个关键行为有 print(f\" ...\") 日志",
        "SPECIAL_MAP 已更新",
        "运行 python demo.py 游戏能正常启动"
    ])

    # ===================== 第三章：关键规范 =====================
    doc.add_page_break()
    add_heading_zh(doc, "第三章  四条铁律（违反必出问题）", level=1)

    add_heading_zh(doc, "铁律一：状态变量必须带卡牌前缀", level=2)
    add_paragraph_zh(doc, "每个卡牌可能需要在单位上存一些临时状态（比如『是否已经触发过』）。这些变量名必须以卡牌缩写开头，否则会和其他卡牌冲突。")
    add_code_block(doc, [
        "# ✅ 正确",
        "minion._songshuqiu_triggered = True",
        "minion._diao_first_attack_done = True",
        "",
        "# ❌ 错误（和别人冲突）",
        "minion._triggered = True",
        "minion._flag = True",
        "minion._tmp = 123"
    ])

    add_heading_zh(doc, "铁律二：回调函数必须用默认参数绑定", level=2)
    add_paragraph_zh(doc, "如果你写了一个『受伤时触发』或『回合结束时触发』的函数，必须这样写：")
    add_code_block(doc, [
        "# ✅ 正确",
        "def on_damage(m=minion, g=game, p=player):",
        "    m.base_keywords['休眠'] = 2",
        "",
        "# ❌ 错误（会导致 bug，指向错误的单位）",
        "def on_damage():",
        "    minion.base_keywords['休眠'] = 2"
    ])
    add_paragraph_zh(doc, "照抄即可，不需要理解原理。")

    add_heading_zh(doc, "铁律三：操作前检查单位是否存活", level=2)
    add_paragraph_zh(doc, "任何时候操作一个单位之前，先问一句：它还活着吗？")
    add_code_block(doc, [
        "# ✅ 正确",
        "if not minion.is_alive():",
        "    return",
        "game.move_minion(minion, new_pos)",
        "",
        "# ❌ 错误（单位可能已经死了）",
        "game.move_minion(minion, new_pos)   # 崩溃！"
    ])

    add_heading_zh(doc, "铁律四：每个关键行为必须打印日志", level=2)
    add_paragraph_zh(doc, "这是为了调试。游戏运行时只有控制台输出，没有图形动画。如果你的效果执行了但没有打印，你根本不知道它有没有触发。")
    add_code_block(doc, [
        "# ✅ 正确",
        'print(f"  {minion.name} 受伤后移动至 {new_pos}")',
        'print(f"  {minion.name} 在 {old_pos} 留下松鼠")',
        "",
        "# ❌ 错误（运行后完全看不到效果是否执行）",
        "game.move_minion(minion, new_pos)   # 静默执行，出了问题找不到"
    ])

    # ===================== 第四章：常见错误 =====================
    doc.add_page_break()
    add_heading_zh(doc, "第四章  常见错误速查表", level=1)

    errors = [
        ("错误 1：装饰器报错 '缺少必要参数'",
         "原因：函数签名漏了 extras 参数。\n解决：在函数参数最后加上 extras=None。",
         ["# ❌ 错误", "def _xxx_special(minion, player, game):",
          "", "# ✅ 正确", "def _xxx_special(minion, player, game, extras=None):"]),

        ("错误 2：效果只在第一次部署时有效",
         "原因：回调函数没有正确绑定单位对象，导致闭包陷阱。\n解决：使用默认参数 m=minion。",
         ["# ✅ 正确", "def on_damage(m=minion, g=game):",
          "    m.base_keywords['休眠'] = 2"]),

        ("错误 3：状态变量和别人冲突",
         "原因：用了 _triggered 这种通用名字。\n解决：状态变量必须以卡牌缩写开头，如 _songshuqiu_triggered。",
         []),

        ("错误 4：SPECIAL_MAP 登记了但游戏找不到效果",
         "原因：SPECIAL_MAP 里的卡牌名和实际注册名不一致（比如漏了括号里的稀有度）。\n解决：复制卡牌在游戏里的完整名称，包括括号。",
         ["# ✅ 正确", '    "松鼠球(铁)": "_songshuqiu_special",',
          "", "# ❌ 错误", '    "松鼠球": "_songshuqiu_special",   # 漏了 (铁)']),

        ("错误 5：导入 effect_utils 失败",
         "原因：项目结构问题或文件不存在。\n解决：检查 card_pools/effect_utils.py 是否存在。",
         []),
    ]

    for title, explain, code in errors:
        add_heading_zh(doc, title, level=2)
        add_paragraph_zh(doc, explain)
        if code:
            add_code_block(doc, code)

    # ===================== 第五章：完整示例 =====================
    doc.add_page_break()
    add_heading_zh(doc, "第五章  完整示例（可以直接抄）", level=1)

    add_paragraph_zh(doc, "下面是一个完全符合规范的完整效果文件。你拿到新卡牌时，照着这个结构填空即可。")
    add_code_block(doc, [
        "# card_pools/underworld_effects.py",
        "",
        "from card_pools.effect_decorator import special",
        "from card_pools.effect_utils import get_adjacent_positions, summon_token",
        "import random",
        "",
        "SPECIAL_MAP = {",
        '    "松鼠球(铁)": "_songshuqiu_special",',
        "    # ... 其他卡牌",
        "}",
        "",
        "",
        "@special",
        "def _songshuqiu_special(minion: \"Minion\", player: \"Player\", game: \"Game\", extras=None) -> None:",
        '    """松鼠球(铁)：受到伤害后向相邻移动一格，在原地留下松鼠。',
        "",
        '    状态变量：',
        '        _songshuqiu_triggered: bool — 标记是否已触发（一次性效果）',
        '    """',
        "    def on_damage(m=minion, g=game, p=player):",
        "        if getattr(m, '_songshuqiu_triggered', False):",
        "            return",
        "        m._songshuqiu_triggered = True",
        "        old_pos = m.position",
        "        candidates = get_adjacent_positions(old_pos, g.board)",
        "        valid = [pos for pos in candidates if pos not in g.board.minion_place]",
        "        if not valid:",
        '            print(f"  {m.name} 受伤但无处可去")',
        "            return",
        "        new_pos = random.choice(valid)",
        "        g.move_minion(m, new_pos)",
        "        summon_token(game=g, name='松鼠', owner=p, position=old_pos, attack=1, health=1)",
        '        print(f"  {m.name} 受伤后移动至 {new_pos}，在 {old_pos} 留下松鼠")',
        "    m._on_take_combat_damage.append(on_damage)",
    ])

    # ===================== 附录 =====================
    doc.add_page_break()
    add_heading_zh(doc, "附录  快捷命令表", level=1)

    add_paragraph_zh(doc, "以下命令在项目根目录的命令行中执行：")
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Light Grid Accent 1'
    hdr = table2.rows[0].cells
    hdr[0].text = '命令'
    hdr[1].text = '作用'
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, "微软雅黑", 11, bold=True)

    cmd_rows = [
        ("python tools/gen_effect.py \"卡牌名(铁)\" \"描述\" --type deploy --file underworld", "生成代码骨架"),
        ("python demo.py", "启动游戏，测试效果"),
        ("python -c \"from card_pools.underworld_effects import SPECIAL_MAP; print(SPECIAL_MAP)\"", "检查 SPECIAL_MAP 是否正确加载"),
        ("git diff", "查看自己修改了哪些文件（提交前检查）"),
    ]
    for cmd, desc in cmd_rows:
        row = table2.add_row().cells
        row[0].text = cmd
        row[1].text = desc
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, "微软雅黑", 10)

    doc.add_paragraph()
    add_paragraph_zh(doc, "— 文档结束 —", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, color=RGBColor(0x88, 0x88, 0x88))

    doc.save(OUTPUT_PATH)
    print(f"文档已保存至: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
